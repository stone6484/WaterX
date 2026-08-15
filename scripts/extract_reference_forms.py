from pathlib import Path
import sys

from docx import Document
from openpyxl import load_workbook


for raw_path in sys.argv[1:]:
    path = Path(raw_path)
    print(f"\n### {path.name}")
    if path.suffix.lower() == ".docx":
        document = Document(path)
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                print(f"P: {text}")
        for table_index, table in enumerate(document.tables):
            print(f"TABLE {table_index}: {len(table.rows)} rows x {len(table.columns)} cols")
            for row in table.rows:
                print(" | ".join(" ".join(cell.text.split()) for cell in row.cells))
    elif path.suffix.lower() == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=False)
        for sheet in workbook.worksheets:
            print(f"SHEET {sheet.title}: {sheet.max_row} rows x {sheet.max_column} cols")
            for row in sheet.iter_rows(min_row=1, max_row=min(sheet.max_row, 120), values_only=True):
                values = ["" if value is None else str(value).replace("\n", " ") for value in row]
                if any(values):
                    print(" | ".join(values))
